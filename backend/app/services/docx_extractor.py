import os
import re
import shutil
import tempfile

import docx

from app.services.extractors.image_summary_service import summarize_image_with_llm


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _save_image_part(img_part, image_dir: str, img_index: int) -> str:
    partname = str(img_part.partname)
    ext = os.path.splitext(partname)[1] if partname else ""
    if not ext:
        ext = ".png"
    img_name = f"img_{img_index}{ext}"
    img_path = os.path.join(image_dir, img_name)
    with open(img_path, "wb") as f:
        f.write(img_part.blob)
    return img_path


def _extract_images_from_paragraphs(paragraphs, rels, image_dir, img_counter):
    image_summaries = []

    for paragraph in paragraphs:
        for run in paragraph.runs:
            blips = run.element.xpath(".//a:blip")
            if not blips:
                continue

            for blip in blips:
                embed_id = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if not embed_id or embed_id not in rels:
                    continue

                img_part = rels[embed_id].target_part
                img_path = _save_image_part(img_part, image_dir, img_counter[0])
                img_counter[0] += 1

                summary_text = summarize_image_with_llm(img_path)
                if summary_text:
                    image_summaries.append(_normalize_text(str(summary_text)))

    return image_summaries


def extract_docx_with_formatting_in_sequence(docx_path: str, image_dir: str | None = None) -> str:
    """
    Extract structured text from DOCX including:
    - headings (converted to markdown #)
    - bullet lists
    - inline images summarized via image_summary_service
    - tables
    - paragraphs in sequential order
    """

    doc = docx.Document(docx_path)
    formatted_output = []
    temp_dir = image_dir or tempfile.mkdtemp(prefix="docx_extract_")

    if image_dir:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir, exist_ok=True)

    rels = doc.part.rels
    para_index = 0
    table_index = 0
    img_counter = [0]

    try:
        for block in doc.element.body.iterchildren():
            if block.tag.endswith("p"):
                paragraph = doc.paragraphs[para_index]
                para_index += 1

                image_texts = _extract_images_from_paragraphs([paragraph], rels, temp_dir, img_counter)
                text = paragraph.text.strip()
                if not text and not image_texts:
                    continue

                style_name = paragraph.style.name.lower()
                if "heading" in style_name:
                    lvl = re.findall(r"\d+", style_name)
                    lvl = int(lvl[0]) if lvl else 1
                    formatted_output.append(f"\n{'#' * lvl} {text}\n")
                elif paragraph._element.xpath(".//w:numPr"):
                    formatted_output.append(f"- {text}")
                else:
                    if text:
                        formatted_output.append(text)

                for image_text in image_texts:
                    formatted_output.append(f"\n### Image\n\n> {image_text}\n")

            elif block.tag.endswith("tbl"):
                table = doc.tables[table_index] if table_index < len(doc.tables) else None
                table_index += 1

                if table is not None:
                    for row in table.rows:
                        cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace("\n", " ")
                            cell_image_texts = _extract_images_from_paragraphs(cell.paragraphs, rels, temp_dir, img_counter)
                            if cell_image_texts:
                                if cell_text:
                                    cell_text += " "
                                cell_text += "Image: " + " / ".join(cell_image_texts)
                            cells.append(cell_text)
                        formatted_output.append("| " + " | ".join(cells) + " |")

                formatted_output.append("\n---\n")

        final_text = "\n".join(formatted_output)
        return re.sub(r"\n{3,}", "\n\n", final_text)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
